import re
import os
import docx
import datetime
import sys
import pandas as pd
import csv
import glob
import subprocess
import arial10
import xlwt
import shutil

getpath = os.path.dirname(os.getcwd())
batfile = os.path.join(getpath, "auto-crt.bat")

def main():
    '''
    Three steps here:
    1. Write and run a bat which will automatically open "SecureCRT" and run the crt Script
    2. Filter some key words
    3. Write the key words to different format, e.g: csv, xls, doc
    The final report will be named as "巡检报告.xls"
    '''

    ##########################################################
    ##Step1: Generate a bat and run it. After you manuall click the "confirm" buttom in the crt, it will quit.
    ##Then we will move to the step2
    ##########################################################

    getbat(batfile)  ##Write a "auto-crt.bat"
    child = subprocess.Popen("auto-crt.bat", cwd=getpath, shell=True, stdin=subprocess.PIPE, stdout=subprocess.PIPE,
                             stderr=subprocess.PIPE)
    stdout, stderr = child.communicate()
    print(child.returncode)

    #########################################################
    #Step2:Handle the text and get some key words
    #########################################################
    for file in gettxtpath():
        host=file.split("\\")[-1].split(".")[0]
        w1 = host + "#show ip ospf neighbor"
        w2 = host + "#show ip bgp summary"
        w3 = host + "#show isis neighbor"
        w4 = host + "#show cable modem summary total"
        w5 = host + "#show envm"
        w6 = host + "#show log"
        Module = []
        Module_Uptime = []
        Power = []
        Temperature = []
        # Startup = []
        # Router = []
        ACL = []
        array = [Power, Temperature]

        ##Filter the neighbor data and the environment data
        with open(file, 'r', encoding='UTF-8') as fp_nei:
            fp_reader = fp_nei.read()
            # fp_reader=fp_reader.replace("\n",'/r/n')
            pat = re.compile(w1 + '(.*?)' + w2, re.S)
            ospfnei = pat.findall(fp_reader)
            pat = re.compile(w2 + '(.*?)' + w3, re.S)
            bgpnei = pat.findall(fp_reader)
            pat = re.compile(w3 + '(.*?)' + w4, re.S)
            isisnei = pat.findall(fp_reader)
            pat=re.compile(w5 + '(.*?)' + w6, re.S)
            envm = pat.findall(fp_reader)
        fp_nei.close()
        neiarray = [ospfnei, bgpnei, isisnei, envm]  ##Get the neighbor list

        ###Change the folder so that the txt.csv can be under Xunjian folder
        path = os.path.join(getpath, 'Xunjian')
        filename=os.path.split(file)[1]   ##Get the filename including the extension
        file1 = filename + ".csv"  ##The file will be like "161.txt.csv"
        file1path = os.path.join(path, file1)

        ###For the remaining data
        with open(file, 'r', encoding='UTF-8') as fp, open(file1path, 'w+', encoding='UTF-8', newline='') as fp1:
            fp1.write("Result\n")
            for line in fp:
                line_stream = line.strip().replace('\x00', '')
                if "System Time" in line:
                    systemtime = line_stream.split(":")[1]
                if "show version" in line:
                    hostname = line.split('#')[0]
                ##########Check the version##############
                if "Running Image" in line:
                    version = line_stream.split(",")[0:3]
                    version = ''.join(version)
                #########Check the product############
                if "Product" in line:
                    product = line.split(':')[1].split(',')[0]
                ##########Check the uptime and the product type#############
                if "System Uptime" in line:
                    uptime = line.split()[2:]
                    uptime = ''.join(uptime)
                    uptime= str(uptime).replace(",",":")
                ############Check the CPU#############################
                if "CPU states" in line:
                    print(line_stream)
                    cpu = line_stream.split(",")[3].strip()
                    cpu = re.findall("[0-9]+\.[0-9]+", cpu)
                ##########Check the Memory############################
                if "MemTotal" in line:
                    MemTotal = re.findall("\d+", line_stream)
                    MemTotal = int(MemTotal[0])
                if "MemFree" in line_stream:
                    MemFree = re.findall("\d+", line_stream)
                    MemFree = int(MemFree[0])
                    ############Check the running state of module#########
                if "Module" in line or "Uptime" in line:
                    # print(line_stream)
                    Module += re.findall("^Module \d+.+$", line_stream)
                    Module_Uptime += re.findall("^Uptime.+$", line_stream)
                if "power module" in line:
                    line_stream = line_stream.strip("[").replace("]", "")
                    Power.append(line_stream)
                if "Module temperature is high" in line or "overheating" in line:
                    Temperature.append(line_stream)
                if "sunrpc" in line:
                    ACL.append(line_stream)
                if "show tech" in line:
                    break
                    ###Check the log###########################
                    # if "failed to become startup" in line:
                    #     line_stream=line_stream.replace(",", " and will")
                    #     #line_stream=line_stream.strip("[").replace("]", "").replace(",", " and will")
                    #     Startup.append(line_stream)
                    # if "-WA-ROUTER-0" in line:
                    #     line_stream=line_stream.replace(",", " and")
                    #     #line_stream = line_stream.strip("[").replace("]", "").replace(",", " and")
                    #     Router.append(line_stream)

     ##########################################################
     ##Step3: Print out the result
     ##########################################################


            print(systemtime)
            print(systemtime,file=fp1)
            print(hostname)
            print(hostname, file=fp1)
            if version.count("Image") <= 2:
                print(version, file=fp1)
            print("The product is %s" % product)
            print(product, file=fp1)
            print(uptime)
            print(uptime, file=fp1)
            if cpu[0] >= str(85):
                print("CPU utilization is normal with %s idle" % cpu[0])
                print("CPU utilization is normal with %s idle" % cpu[0], file=fp1)
            else:
                print("CPU utilization is abnormal with %s idle" % cpu[0], file=fp1)
            if MemFree / MemTotal > 0.45:
                print("Memory is with good with %.2f%%" % (MemFree / MemTotal * 100), file=fp1)
            else:
                print("Memory is with abnormal with %.2f%%" % (MemFree / MemTotal * 100), file=fp1)

           ##Get the list of "Module" and "Uptime".
            c = []
            print(Module)
            print(Module_Uptime)
            for i in range(len(Module)):
                if Module_Uptime != []:
                    c.append(Module[i] + "\n")
                    c.append(Module_Uptime[i] + "\n")
                else:
                    c.append(Module[i] + "\n")
            c_1 = ''.join(c)
            Module_F = [c_1]
            print(Module_F)

            ##Manual Write to the csv as the normal print function can't perform well for the list type.
            #############################################################
            w = csv.writer(fp1)
            ##Write the module and uptime information
            w.writerow(Module_F)
            ##Write the neighbor list including ospf, bgp, isis
            for i in range(len(neiarray)):
                tip=['Not configured']
                if neiarray[i]!=['\n']:
                    w.writerow(neiarray[i])
                else:
                    w.writerow(tip)


            ##For the other logs like Power and temperature....
            for id in range(len(array)):
                if array[id]==[]:
                    print("Good")
                    print("Good",file=fp1)
                else:
                    #print(array[id])
                    out=','.join(array[id])
                    print(out)
                    print(out, file=fp1)

            ##For ACL ###########################################
            if ACL==[]:
                print("No sunrpc ACL is detected. Please implement the ACL",file=fp1)
            else:
                print(ACL,file=fp1)

        fp.close()
        fp1.close()

    GetFormated() ##Get the Final output of Xunjian-Report.csv
    Getxlsx()   ##Get the Xunjian-Report.xls
    # GetDocx()    ##Get the docx format




def getbat(batfile): ##Get an auto-crt.bat
    # getpath=os.path.dirname(os.getcwd())
    # #batfile=os.path.join(getpath,"auto-crt.bat")
    with open(batfile,'w+') as batman:
        batman.write("@echo off")
        batman.write("\n")
        batman.write("start {path}\App\SecureCRT\SecureCRT.exe /SCRIPT {path}\Xunjian-crt.py".format(path=getpath))
        batman.write("\n")
        batman.write(":end")
        batman.write("\n")
        batman.write("exit")
    batman.close()

def gettxtpath():  ##'.txt'.
    #Get all the filename with txt format. It is for data filtering
    ##The filename is under the Log folder
    path=os.path.join(getpath, 'Log')
    file_path=[]
    f_list=os.listdir(path)
    for file in f_list:
        #os.path.splitext()
        if os.path.splitext(file)[1] == '.txt':
            filepath=os.path.join(path,file)
            file_path.append(filepath)
    return file_path



def getcsvName():
    #Get all the filename with txt.csv format..
    path = os.getcwd()
    file_list=[]
    csv_list=[]
    f_list=os.listdir(path)
    for file in f_list:
        #os.path.splitext()
        if os.path.splitext(file)[1] == '.csv':
            file_list.append(file)
    for file in file_list:
        if file.endswith(".txt.csv"):   ###It has contained the format
            csv_list.append(file)
    return csv_list    ##It has a list for the ".txt.csv" file like "161.txt.csv"

def CsvMerge(file1, file2,output):
    ##It is for preparation of the getFormated
    ##Merge the "Format.csv" and the "text.csv.csv
    df1=pd.read_csv(file1)
    df2=pd.read_table(file2)
    df3=pd.concat([df1,df2],axis=1)
    print(df3)
    with open(output,'w+',encoding="utf-8") as f:
        df3.to_csv(f,index=False)
    f.close()

def getResult():
    ##Get the Final output of Xunjian-Report.csv. Merged all the csv with txt.csv.csv
    path = os.getcwd()
    allfiles = glob.glob(path + "/*.txt.csv.csv")
    df_out_filename='Xunjian-Report.csv'
    #write_headers=True
    with open(df_out_filename,'w+',newline='') as fout:
        writer=csv.writer(fout)
        for filename in allfiles:
            with open(filename) as fin:
                reader=csv.reader(fin)
                #headers=next(reader)
                #if write_headers:
                #    write_headers=False     # Only write headers once.
                    #writer.writerow(headers)
                writer.writerows(reader)
                writer.writerow("\n")   ##Insert blank row at the end

def Getxlsx():
    '''
    Read the csv and save it with xls format
    Also Setting the format of the cells
    ##Wrap at right
    ##Horz_left
    ##Vert_center
    '''
    myexcel=xlwt.Workbook()
    mysheet=FitSheetWrapper(myexcel.add_sheet("Xunjian"))  ##Apply the
    ##Set the format of the cells
    al1=xlwt.Alignment()
    al1.wrap=xlwt.Alignment.WRAP_AT_RIGHT
    al1.horz=xlwt.Alignment.HORZ_LEFT
    al1.vert=xlwt.Alignment.VERT_CENTER
    style=xlwt.XFStyle()
    style.alignment=al1
    ###Read the csv and save it with xls format
    csvfile=open("Xunjian-Report.csv","r")
    reader=csv.reader(csvfile)
    l=0
    for line in reader:
        r=0
        for i in line:
            mysheet.write(l,r,i,style)
            r=r+1
        l=l+1
    myexcel.save("FinalReport.xls")
    ###Copy the file to upper folder which is is in the same path of the crt.py
    shutil.copyfile("FinalReport.xls","{path}/FinalReport.xls".format(path=getpath))

class FitSheetWrapper(object):
    #https://stackoverflow.com/questions/6929115/python-xlwt-accessing-existing-cell-content-auto-adjust-column-width
    """Try to fit columns to max size of any entry.
    To use, wrap this around a worksheet returned from the
    workbook's add_sheet method, like follows:
        sheet = FitSheetWrapper(book.add_sheet(sheet_name))
    The worksheet interface remains the same: this is a drop-in wrapper
    for auto-sizing columns.
    """
    def __init__(self, sheet):
        self.sheet = sheet
        self.widths = dict()

    def write(self, r, c, label='', *args, **kwargs):
        self.sheet.write(r, c, label, *args, **kwargs)
        width = int(arial10.fitwidth(label))

        if width > self.widths.get(c, 0):
            self.widths[c] = width
            self.sheet.col(c).width = width

    def __getattr__(self, attr):
        return getattr(self.sheet, attr)

def GetFormated():
    for file in getcsvName():
        file1=file+'.csv'
        CsvMerge('Format.csv',file, file1)###Output will be like 161.txt.csv.csv
        getResult()  ##Get the Xunjian-Report.csv


def GetDocx():   ##Get docx  ##It requires the csv doesn't have blank rows
    csvfile = 'Xunjian-Report.csv'
    doc = docx.Document()
    date = datetime.datetime.now()
    with open(csvfile, newline='') as f:
        csv_reader = csv.reader(f)
        csv_headers = next(csv_reader)
        csv_cols = len(csv_headers)
        # Adding Subject of the document and the date
        doc.add_heading('巡检报告', level=0)
        doc.add_heading('编辑日期: %s/%s/%s' % (date.day, date.month, date.year), level=1)

        table = doc.add_table(rows=1, cols=csv_cols)
        hdr_cells = table.rows[0].cells
        ##表格的标题
        for i in range(csv_cols):
            hdr_cells[i].text = csv_headers[i]
            ##每行的内容
        for row in csv_reader:
            row_cells = table.add_row().cells
            for j in range(csv_cols):
                row_cells[j].text = row[j]
    doc.add_page_break()
    ##粗体
    # paragraph.add_run('xxx').bold=True
    doc.save("Xunjian-Report.docx")

if __name__ == "__main__":
    main()


